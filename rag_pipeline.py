
# import os
# import hashlib
# import PyPDF2
# from docx import Document
# from datetime import datetime
# import logging

# # Configure logging
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
# logger = logging.getLogger(__name__)

# # --- LangChain modular imports (Oct 2025+) ---
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores.pgvector import PGVector
# from langchain_openai import ChatOpenAI
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.runnables import RunnableLambda
# # ---------------------------------------------
# from config import RAGConfig,rag_config 

# # # ============ CONFIGURATION ============
# from dotenv import load_dotenv
# load_dotenv()
# CONNECTION_STRING = (
#     f"postgresql+psycopg2://{os.getenv('PGUSER')}:{os.getenv('PGPASSWORD')}"
#     f"@{os.getenv('PGHOST')}:{os.getenv('PGPORT')}/{os.getenv('PGDATABASE')}"
# )
# if not all([
#     os.getenv('PGUSER'),
#     os.getenv('PGPASSWORD'),
#     os.getenv('PGHOST'),
#     os.getenv('PGPORT'),
#     os.getenv('PGDATABASE')
# ]):
#     raise ValueError("One or more required Postgres environment variables are missing or empty.")

# # Cache for embedding model (singleton pattern)
# _embedding_cache = None
# # =================main function for embeddings================
# def get_embeddings():
#     """Get embedding model based on environment variable."""
#     global _embedding_cache
#     try: 
#         logger.info(f"RAG Config in get_embeddings: model_name={rag_config.model_name}, model_type={rag_config.model_type}, local_embedding_name={rag_config.local_embedding_name}")    
#         # Log function call with parameters
#         logger.info("get_embeddings called")
#         logger.info(f"model_type from config: {rag_config.model_type}")
        
#         if _embedding_cache is not None:
#             logger.info("Using cached embedding model")
#             return _embedding_cache

#         if rag_config.model_type.lower() == "openai":
#             logger.info("Loading OpenAI embeddings...")
#             logger.info(f"OpenAI API Key present: {'Yes' if rag_config.openai_api_key else 'No'}")
#             logger.info(f"**********model_name: {rag_config.model_name},model_type: {rag_config.model_type}, openai_api_key: {rag_config.openai_api_key}, self.document_collection_name: {rag_config.document_collection_name} , rag_config.document_collection_path: {rag_config.document_collection_path} **********")

#             from langchain_openai import OpenAIEmbeddings
#             logger.info(f"starting to load openai embeddings with api key: {rag_config.openai_api_key}")

#             _embedding_cache = OpenAIEmbeddings(api_key=rag_config.openai_api_key)
#             logger.info("OpenAI embeddings loaded successfully")
#         else:
#             logger.info("Loading local HuggingFace embeddings (this may take a moment)...")
#             logger.info(f"LOCAL_EMBEDDING_NAME from env: {rag_config.local_embedding_name}")
#             from langchain_huggingface import HuggingFaceEmbeddings

#             logger.info(f"Loading model: {rag_config.local_embedding_name}")
#             _embedding_cache = HuggingFaceEmbeddings(
#                 model_name=rag_config.local_embedding_name,
#                 model_kwargs={'device': 'cpu'},
#                 encode_kwargs={'normalize_embeddings': True}
#             )
#             logger.info("HuggingFace embeddings loaded successfully!")
#     except Exception as e:
#         logger.error(f"Error accessing rag_config: {e}")
    
    
#     return _embedding_cache


# def file_hash(path: str) -> str:
#     """Generate a stable hash of a file for change detection."""
#     with open(path, "rb") as f:
#         return hashlib.md5(f.read()).hexdigest()


# def clean_text(text: str) -> str:
#     """Remove NUL characters and other problematic characters from text."""
#     # Remove NUL characters
#     text = text.replace('\x00', '')
#     # Remove other control characters except newlines and tabs
#     text = ''.join(char for char in text if char == '\n' or char == '\t' or ord(char) >= 32)
#     return text.strip()


# def load_documents(folder_path: str):
#     """Read all supported files from a folder."""
#     docs = []
#     for file in os.listdir(folder_path):
#         path = os.path.join(folder_path, file)
#         if not os.path.isfile(path):
#             continue
#         if not file.lower().endswith((".pdf", ".txt", ".docx")):
#             continue

#         try:
#             if file.endswith(".pdf"):
#                 with open(path, "rb") as f:
#                     reader = PyPDF2.PdfReader(f)
#                     text = "".join([page.extract_text() or "" for page in reader.pages])
#             elif file.endswith(".txt"):
#                 with open(path, "r", encoding="utf-8", errors='ignore') as f:
#                     text = f.read()
#             elif file.endswith(".docx"):
#                 doc = Document(path)
#                 text = "\n".join([p.text for p in doc.paragraphs])
#             else:
#                 continue

#             # Clean the text to remove problematic characters
#             text = clean_text(text)
            
#             if not text:
#                 print(f"Skipping empty document: {file}")
#                 continue

#             docs.append({
#                 "text": text,
#                 "metadata": {
#                     "filename": file,
#                     "path": path,
#                     "hash": file_hash(path),
#                     "timestamp": datetime.fromtimestamp(os.path.getmtime(path)).isoformat(),
#                 }
#             })
#         except Exception as e:
#             print(f"Error processing {file}: {str(e)}")
#             continue
    
#     return docs


# def chunk_texts(docs, chunk_size=1000, overlap=200):
#     """Split documents into overlapping text chunks."""
#     splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
#     chunks = []
#     for d in docs:
#         parts = splitter.create_documents([d["text"]], metadatas=[d["metadata"]])
#         chunks.extend(parts)
#     return chunks


# def build_pgvector_store(chunks):
#     """Embed and store document chunks in pgvector."""
#     embeddings = get_embeddings()
#     store = PGVector.from_documents(
#         documents=chunks,
#         embedding=embeddings,
#         connection_string=CONNECTION_STRING,
#         collection_name=rag_config.document_collection_name,
#     )
#     return store


# def load_existing_pgvector_store():
#     """Connect to existing pgvector collection."""
#     embeddings = get_embeddings()
#     logger.info(f"Loading existing pgvector store: {rag_config.document_collection_name}")
#     return PGVector(
#         connection_string=CONNECTION_STRING,
#         embedding_function=embeddings,
#         collection_name=rag_config.document_collection_name,
#     )


# def get_stored_hashes(store):
#     """Retrieve already-indexed file hashes."""
#     try:
#         with store._conn.cursor() as cur:
#             cur.execute(f"SELECT metadata->>'hash' FROM {rag_config.document_collection_name};")
#             return {r[0] for r in cur.fetchall() if r[0]}
#     except Exception:
#         return set()


# def update_index(force=False):
#     """Re-index new or changed files."""
#     if not force and not os.path.exists(rag_config.document_collection_path):
#         print(f"Folder {rag_config.document_collection_path} not found")
#         return
    
#     existing_store = load_existing_pgvector_store()
#     existing_hashes = get_stored_hashes(existing_store)

#     new_chunks = []
#     for doc in load_documents(rag_config.document_collection_path):
#         if doc["metadata"]["hash"] not in existing_hashes:
#             new_chunks.extend(chunk_texts([doc]))

#     if new_chunks:
#         print(f"Indexing {len(new_chunks)} new/updated chunks...")
#         build_pgvector_store(new_chunks)
#     else:
#         print("No new or changed documents.")


# def format_docs(docs):
#     """Format retrieved documents for the prompt."""
#     return "\n\n".join(doc.page_content for doc in docs)

# #main function to create the QA chain
# def make_qa_chain(skip_update=True):
#     """Create a Retrieval-Augmented QA chain using LCEL.
    
#     Args:
#         skip_update: If True, skip checking for new documents (faster).
#         config: RAGConfig instance or None to use global config
#         **kwargs: Override individual config parameters
#     """
#     # Use provided config or global config
#     # if config is None:
#     #     config = rag_config
    
#     # Override with any kwargs
#     #if kwargs:
#         # # Create a copy to avoid modifying the original
#         # config = RAGConfig()
#         # config.model_name = kwargs.get('model_name', rag_config.model_name)
#         # config.model_type = kwargs.get('model_type', rag_config.model_type)
#         # config.system_prompt = kwargs.get('system_prompt', rag_config.system_prompt)
#         # config.document_collection = kwargs.get('document_collection', rag_config.document_collection)
#     if not skip_update:
#      update_index()

#     store = load_existing_pgvector_store()
#     logger.info(f"Loaded PGVector store: {rag_config.document_collection_name}")
#     retriever = store.as_retriever(search_kwargs={"k": 5})
#     logger.info("Retriever created successfully")
    
#     # Create LLM based on model type
#     if rag_config.model_type.lower() == "openai":
#         logger.info("Using OpenAI LLM")
#         logger.info(f"model_name: {rag_config.model_name}")
#         try:
#             logger.info(f"OpenAI API Key present: {'Yes' if rag_config.openai_api_key else 'No'}")
#             llm = ChatOpenAI(model=rag_config.model_name, temperature=0, api_key=rag_config.openai_api_key)
#             logger.info("ChatOpenAI initialized successfully (make_qa_chain)")
#         except Exception as e:
#             logger.error(f"Error initializing ChatOpenAI: {e}")
#             raise
#     elif rag_config.model_type.lower() == "local":
#         try:
#             # For local models, you can use Ollama or other local LLM providers
#             from langchain_community.llms import Ollama
#             llm = Ollama(model=rag_config.model_name, temperature=0)
#             print(f"Using local Ollama model: {rag_config.model_name}")
#         except ImportError: 
#             print("Ollama not available, falling back to OpenAI")
#     elif rag_config.model_type.lower() == "huggingface":
#         try:
#             # For HuggingFace models
#             from langchain_huggingface import HuggingFacePipeline
#             llm = HuggingFacePipeline.from_model_id(
#                 model_id=rag_config.model_name,
#                 task="text-generation",
#                 model_kwargs={"temperature": 0, "max_length": 512}
#             )
#             print(f"Using HuggingFace model: {rag_config.model_name}")
#         except ImportError:
#             print("HuggingFace not available, falling back to OpenAI")
#     else:
#         print(f"Unknown model type '{rag_config.model_type}', defaulting to OpenAI")
#         llm = ChatOpenAI(model=rag_config.model_name, temperature=0, api_key=rag_config.openai_api_key)

#     # Create prompt template
#     template = """{system_prompt}

# {context}

# Question: {input}

# Answer:"""
    
#     prompt = ChatPromptTemplate.from_template(template)

#     # Build the chain using LCEL (LangChain Expression Language)
#     def get_input(x):
#         """Extract input from dict or return string as-is."""
#         if isinstance(x, dict):
#             return x.get("input", x.get("question", ""))
#         return x

#     # Create the base chain
#     base_chain = (
#         {
#             # "context": lambda x: format_docs(retriever.invoke(get_input(x))),
#             "context": lambda x: format_docs(retriever.invoke(get_input(x))),
#             "input": get_input,
#             "system_prompt": lambda x: rag_config.system_prompt
#         }
#         | prompt
#         | llm
#         | StrOutputParser()
#     )
    
#     # Wrap to return dictionary format for compatibility
#     def wrapped_chain(x):
#         result = base_chain.invoke(x)
#         return {"output_text": result, "answer": result}
    
#     # Create a runnable wrapper
#     qa_chain = RunnableLambda(wrapped_chain)
    
#     return qa_chain







import os
import hashlib
import PyPDF2
from docx import Document
from datetime import datetime
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- LangChain modular imports (Oct 2025+) ---
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores.pgvector import PGVector
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnableLambda
# ---------------------------------------------
from config import RAGConfig,rag_config 

# # ============ CONFIGURATION ============
from dotenv import load_dotenv
load_dotenv()
CONNECTION_STRING = (
    f"postgresql+psycopg2://{os.getenv('PGUSER')}:{os.getenv('PGPASSWORD')}"
    f"@{os.getenv('PGHOST')}:{os.getenv('PGPORT')}/{os.getenv('PGDATABASE')}"
)
if not all([
    os.getenv('PGUSER'),
    os.getenv('PGPASSWORD'),
    os.getenv('PGHOST'),
    os.getenv('PGPORT'),
    os.getenv('PGDATABASE')
]):
    raise ValueError("One or more required Postgres environment variables are missing or empty.")

# Cache for embedding model (singleton pattern)
_embedding_cache = None
# =================main function for embeddings================
def get_embeddings():
    """Get embedding model based on environment variable."""
    global _embedding_cache
    
    # Log function call with parameters
    logger.info("get_embeddings called")
    logger.info(f"model_type from config: {rag_config.model_type}")
    
    if _embedding_cache is not None:
        logger.info("Using cached embedding model")
        return _embedding_cache

    if rag_config.model_type.lower() == "openai":
        logger.info("Loading OpenAI embeddings...")
        logger.info(f"OpenAI API Key present: {'Yes' if rag_config.openai_api_key else 'No'}")
        logger.info(f"**********model_name: {rag_config.model_name},model_type: {rag_config.model_type}, openai_api_key: {rag_config.openai_api_key}, self.document_collection_name: {rag_config.document_collection_name} , rag_config.document_collection_path: {rag_config.document_collection_path} **********")

        from langchain_openai import OpenAIEmbeddings
        logger.info(f"starting to load openai embeddings with api key: {rag_config.openai_api_key}")

        _embedding_cache = OpenAIEmbeddings(api_key=rag_config.openai_api_key)
        logger.info("OpenAI embeddings loaded successfully")
    else:
        logger.info("Loading local HuggingFace embeddings (this may take a moment)...")
        logger.info(f"LOCAL_EMBEDDING_NAME from env: {rag_config.local_embedding_name}")
        from langchain_huggingface import HuggingFaceEmbeddings

        logger.info(f"Loading model: {rag_config.local_embedding_name}")
        _embedding_cache = HuggingFaceEmbeddings(
            model_name=rag_config.local_embedding_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        logger.info("HuggingFace embeddings loaded successfully!")
    
    return _embedding_cache


def file_hash(path: str) -> str:
    """Generate a stable hash of a file for change detection."""
    with open(path, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()


def clean_text(text: str) -> str:
    """Remove NUL characters and other problematic characters from text."""
    # Remove NUL characters
    text = text.replace('\x00', '')
    # Remove other control characters except newlines and tabs
    text = ''.join(char for char in text if char == '\n' or char == '\t' or ord(char) >= 32)
    return text.strip()


def load_documents(folder_path: str):
    """Read all supported files from a folder."""
    docs = []
    for file in os.listdir(folder_path):
        path = os.path.join(folder_path, file)
        if not os.path.isfile(path):
            continue
        if not file.lower().endswith((".pdf", ".txt", ".docx")):
            continue

        try:
            if file.endswith(".pdf"):
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = "".join([page.extract_text() or "" for page in reader.pages])
            elif file.endswith(".txt"):
                with open(path, "r", encoding="utf-8", errors='ignore') as f:
                    text = f.read()
            elif file.endswith(".docx"):
                doc = Document(path)
                text = "\n".join([p.text for p in doc.paragraphs])
            else:
                continue

            # Clean the text to remove problematic characters
            text = clean_text(text)
            
            if not text:
                print(f"Skipping empty document: {file}")
                continue

            docs.append({
                "text": text,
                "metadata": {
                    "filename": file,
                    "path": path,
                    "hash": file_hash(path),
                    "timestamp": datetime.fromtimestamp(os.path.getmtime(path)).isoformat(),
                }
            })
        except Exception as e:
            print(f"Error processing {file}: {str(e)}")
            continue
    
    return docs


def chunk_texts(docs, chunk_size=1000, overlap=200):
    """Split documents into overlapping text chunks."""
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
    chunks = []
    for d in docs:
        parts = splitter.create_documents([d["text"]], metadatas=[d["metadata"]])
        chunks.extend(parts)
    return chunks


def build_pgvector_store(chunks):
    """Embed and store document chunks in pgvector."""
    embeddings = get_embeddings()
    store = PGVector.from_documents(
        documents=chunks,
        embedding=embeddings,
        connection_string=CONNECTION_STRING,
        collection_name=rag_config.document_collection_name,
    )
    return store


def load_existing_pgvector_store():
    """Connect to existing pgvector collection."""
    embeddings = get_embeddings()
    logger.info(f"Loading existing pgvector store: {rag_config.document_collection_name}")
    return PGVector(
        connection_string=CONNECTION_STRING,
        embedding_function=embeddings,
        collection_name=rag_config.document_collection_name,
    )


def get_stored_hashes(store):
    """Retrieve already-indexed file hashes."""
    try:
        with store._conn.cursor() as cur:
            cur.execute(f"SELECT metadata->>'hash' FROM {rag_config.document_collection_name};")
            return {r[0] for r in cur.fetchall() if r[0]}
    except Exception:
        return set()


def update_index(force=False):
    """Re-index new or changed files."""
    if not force and not os.path.exists(rag_config.document_collection_path):
        print(f"Folder {rag_config.document_collection_path} not found")
        return
    
    existing_store = load_existing_pgvector_store()
    existing_hashes = get_stored_hashes(existing_store)

    new_chunks = []
    for doc in load_documents(rag_config.document_collection_path):
        if doc["metadata"]["hash"] not in existing_hashes:
            new_chunks.extend(chunk_texts([doc]))

    if new_chunks:
        print(f"Indexing {len(new_chunks)} new/updated chunks...")
        build_pgvector_store(new_chunks)
    else:
        print("No new or changed documents.")


def format_docs(docs):
    """Format retrieved documents for the prompt."""
    return "\n\n".join(doc.page_content for doc in docs)

#main function to create the QA chain
def make_qa_chain(skip_update=True):
    """Create a Retrieval-Augmented QA chain using LCEL.
    
    Args:
        skip_update: If True, skip checking for new documents (faster).
        config: RAGConfig instance or None to use global config
        **kwargs: Override individual config parameters
    """
    # Use provided config or global config
    # if config is None:
    #     config = rag_config
    
    # Override with any kwargs
    #if kwargs:
        # # Create a copy to avoid modifying the original
        # config = RAGConfig()
        # config.model_name = kwargs.get('model_name', rag_config.model_name)
        # config.model_type = kwargs.get('model_type', rag_config.model_type)
        # config.system_prompt = kwargs.get('system_prompt', rag_config.system_prompt)
        # config.document_collection = kwargs.get('document_collection', rag_config.document_collection)
    if not skip_update:
     update_index()

    store = load_existing_pgvector_store()
    logger.info(f"Loaded PGVector store: {rag_config.document_collection_name}")
    retriever = store.as_retriever(search_kwargs={"k": 5})
    logger.info("Retriever created successfully")
    # Create LLM based on model type
    if rag_config.model_type.lower() == "openai":
        logger.info("Using OpenAI LLM")
        logger.info(f"model_name: {rag_config.model_name}")
        llm = ChatOpenAI(model=rag_config.model_name, temperature=0, api_key=rag_config.openai_api_key)
    elif rag_config.model_type.lower() == "local":
        try:
            # For local models, you can use Ollama or other local LLM providers
            from langchain_community.llms import Ollama
            llm = Ollama(model=rag_config.model_name, temperature=0)
            print(f"Using local Ollama model: {rag_config.model_name}")
        except ImportError: 
            print("Ollama not available, falling back to OpenAI")
    elif rag_config.model_type.lower() == "huggingface":
        try:
            # For HuggingFace models
            from langchain_huggingface import HuggingFacePipeline
            llm = HuggingFacePipeline.from_model_id(
                model_id=rag_config.model_name,
                task="text-generation",
                model_kwargs={"temperature": 0, "max_length": 512}
            )
            print(f"Using HuggingFace model: {rag_config.model_name}")
        except ImportError:
            print("HuggingFace not available, falling back to OpenAI")
    else:
        print(f"Unknown model type '{rag_config.model_type}', defaulting to OpenAI")
        llm = ChatOpenAI(model=rag_config.model_name, temperature=0, api_key=rag_config.openai_api_key)

    # Create prompt template for business requirement analysis
    template = """You are a Business Analysis AI assistant specialized in helping users define or analyze business systems and create Business Requirements Documents (BRDs). Your goals: (1) When the user describes a system or use case, first check if a similar system exists in the RAG database using semantic similarity search. If a similar system is found, summarize it briefly and say: 'A similar system exists called <SystemName>. Would you like to review or extend its BRD?' Then provide the summary of the existing system and its main features. (2) If no similar system exists, start guiding the user to create a new BRD by asking structured questions, such as: What is the purpose of the system? Who are the target users or stakeholders? What are the main features or functional requirements? What are the non-functional requirements (e.g., performance, security)? What integrations or dependencies does it have? What are the expected outputs or success metrics? (3) Keep the conversation contextual: store user answers and progressively build a structured BRD outline. Use professional BRD formatting with sections: Overview, Objectives, Functional Requirements, Non-Functional Requirements, Integration, Dependencies, Timeline, Risks, etc. (4) When the BRD is ready, present it as a structured draft and ask: 'Would you like to export this BRD as a document (Word, PDF, or Markdown)?' Maintain a professional, concise, and proactive tone.

Based on the following context documents:

{context}

Question: {input}

Answer:"""
    
    prompt = ChatPromptTemplate.from_template(template)

    # Build the chain using LCEL (LangChain Expression Language)
    def get_input(x):
        """Extract input from dict or return string as-is."""
        if isinstance(x, dict):
            return x.get("input", x.get("question", ""))
        return x

    # Create the base chain
    def get_context_with_business_analysis(x):
        """Retrieve context with business requirement analysis and feature gap detection."""
        query = get_input(x)
        query_lower = query.lower()
        
        # Extract explicit filter criteria if provided
        if isinstance(x, dict) and "filter" in x:
            docs = retriever.invoke(query, filter=x["filter"])
        elif isinstance(x, dict) and "document_type" in x:
            filter_criteria = {"filename": {"$regex": f".*\\.{x['document_type']}$"}}
            docs = retriever.invoke(query, filter=filter_criteria)
        else:
            # Business requirement analysis - detect intent and context
            business_intent_keywords = {
                # Business requirement indicators
                "need": "business_requirement",
                "require": "business_requirement", 
                "want": "business_requirement",
                "should": "business_requirement",
                "must": "business_requirement",
                "feature": "feature_request",
                "functionality": "feature_request",
                "capability": "feature_request",
                "implement": "implementation",
                "develop": "implementation",
                "build": "implementation",
                # Current state analysis indicators
                "current": "current_state",
                "existing": "current_state",
                "already": "current_state",
                "have": "current_state",
                "does": "current_state",
                "can": "current_state"
            }
            
            # Context-specific document filters
            context_filters = {
                "brd": {"path": {"$regex": ".*brd.*"}},
                "business requirement": {"path": {"$regex": ".*brd.*"}},
                "requirements": {"path": {"$regex": ".*brd.*"}},
                "policy": {"path": {"$regex": ".*polic.*"}},
                "policies": {"path": {"$regex": ".*polic.*"}},
                "procedure": {"path": {"$regex": ".*procedure.*"}},
                "procedures": {"path": {"$regex": ".*procedure.*"}},
                "technical": {"path": {"$regex": ".*tech.*"}},
                "system": {"path": {"$regex": ".*(system|arch|design).*"}},
                "architecture": {"path": {"$regex": ".*(system|arch|design).*"}},
                "manual": {"path": {"$regex": ".*manual.*"}},
                "guide": {"path": {"$regex": ".*guide.*"}},
                "specification": {"path": {"$regex": ".*spec.*"}},
                "functional": {"path": {"$regex": ".*(functional|spec).*"}},
            }
            
            # Detect business intent
            detected_intent = None
            for keyword, intent in business_intent_keywords.items():
                if keyword in query_lower:
                    detected_intent = intent
                    break
            
            # Detect document context
            detected_context = None
            detected_filter = None
            for keyword, filter_criteria in context_filters.items():
                if keyword in query_lower:
                    detected_context = keyword
                    detected_filter = filter_criteria
                    logger.info(f"Auto-detected document context: '{keyword}'")
                    break
            
            # Smart retrieval strategy based on business intent
            if detected_intent == "business_requirement":
                # For new business requirements, search BRDs first, then system docs
                if detected_filter:
                    docs = retriever.invoke(query, filter=detected_filter)
                else:
                    # Search BRDs and system docs for gap analysis
                    brd_docs = retriever.invoke(query, filter={"path": {"$regex": ".*brd.*"}})
                    system_docs = retriever.invoke(query, filter={"path": {"$regex": ".*(system|tech|spec).*"}})
                    docs = brd_docs + system_docs
                    docs = docs[:5]  # Limit to top 5 results
                    
            elif detected_intent == "current_state":
                # For current state queries, prioritize system and technical docs
                if detected_filter:
                    docs = retriever.invoke(query, filter=detected_filter)
                else:
                    docs = retriever.invoke(query, filter={"path": {"$regex": ".*(system|tech|manual|guide).*"}})
                    
            elif detected_filter:
                # Use detected context filter
                docs = retriever.invoke(query, filter=detected_filter)
            else:
                # Default comprehensive search
                docs = retriever.invoke(query)
            
            # Log the analysis for debugging
            logger.info(f"Business Intent: {detected_intent}, Context: {detected_context}")
        
        return format_docs(docs)
    
    base_chain = (
        {
            "context": get_context_with_business_analysis,
            "input": get_input
        }
        | prompt
        | llm
        | StrOutputParser()
    )
    
    # Wrap to return dictionary format for compatibility
    def wrapped_chain(x):
        result = base_chain.invoke(x)
        return {"output_text": result, "answer": result}
    
    # Create a runnable wrapper
    qa_chain = RunnableLambda(wrapped_chain)
    
    return qa_chain

# //works fine
# import os
# import hashlib
# import PyPDF2
# from docx import Document
# from datetime import datetime
# import logging

# # Configure logging
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
# logger = logging.getLogger(__name__)

# # --- LangChain modular imports (Oct 2025+) ---
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores.pgvector import PGVector
# from langchain_openai import ChatOpenAI
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.runnables import RunnableLambda
# # ---------------------------------------------
# from config import RAGConfig,rag_config 

# # # ============ CONFIGURATION ============
# from dotenv import load_dotenv
# load_dotenv()
# CONNECTION_STRING = (
#     f"postgresql+psycopg2://{os.getenv('PGUSER')}:{os.getenv('PGPASSWORD')}"
#     f"@{os.getenv('PGHOST')}:{os.getenv('PGPORT')}/{os.getenv('PGDATABASE')}"
# )
# if not all([
#     os.getenv('PGUSER'),
#     os.getenv('PGPASSWORD'),
#     os.getenv('PGHOST'),
#     os.getenv('PGPORT'),
#     os.getenv('PGDATABASE')
# ]):
#     raise ValueError("One or more required Postgres environment variables are missing or empty.")

# # Cache for embedding model (singleton pattern)
# _embedding_cache = None
# # =================main function for embeddings================
# def get_embeddings():
#     """Get embedding model based on environment variable."""
#     global _embedding_cache
    
#     # Log function call with parameters
#     logger.info("get_embeddings called")
#     logger.info(f"model_type from config: {rag_config.model_type}")
    
#     if _embedding_cache is not None:
#         logger.info("Using cached embedding model")
#         return _embedding_cache

#     if rag_config.model_type.lower() == "openai":
#         logger.info("Loading OpenAI embeddings...")
#         logger.info(f"OpenAI API Key present: {'Yes' if rag_config.openai_api_key else 'No'}")
#         logger.info(f"**********model_name: {rag_config.model_name},model_type: {rag_config.model_type}, openai_api_key: {rag_config.openai_api_key}, self.document_collection_name: {rag_config.document_collection_name} , rag_config.document_collection_path: {rag_config.document_collection_path} **********")

#         from langchain_openai import OpenAIEmbeddings
#         logger.info(f"starting to load openai embeddings with api key: {rag_config.openai_api_key}")

#         _embedding_cache = OpenAIEmbeddings(api_key=rag_config.openai_api_key)
#         logger.info("OpenAI embeddings loaded successfully")
#     else:
#         logger.info("Loading local HuggingFace embeddings (this may take a moment)...")
#         logger.info(f"LOCAL_EMBEDDING_NAME from env: {rag_config.local_embedding_name}")
#         from langchain_huggingface import HuggingFaceEmbeddings

#         logger.info(f"Loading model: {rag_config.local_embedding_name}")
#         _embedding_cache = HuggingFaceEmbeddings(
#             model_name=rag_config.local_embedding_name,
#             model_kwargs={'device': 'cpu'},
#             encode_kwargs={'normalize_embeddings': True}
#         )
#         logger.info("HuggingFace embeddings loaded successfully!")
    
#     return _embedding_cache


# def file_hash(path: str) -> str:
#     """Generate a stable hash of a file for change detection."""
#     with open(path, "rb") as f:
#         return hashlib.md5(f.read()).hexdigest()


# def clean_text(text: str) -> str:
#     """Remove NUL characters and other problematic characters from text."""
#     # Remove NUL characters
#     text = text.replace('\x00', '')
#     # Remove other control characters except newlines and tabs
#     text = ''.join(char for char in text if char == '\n' or char == '\t' or ord(char) >= 32)
#     return text.strip()


# def load_documents(folder_path: str):
#     """Read all supported files from a folder."""
#     docs = []
#     for file in os.listdir(folder_path):
#         path = os.path.join(folder_path, file)
#         if not os.path.isfile(path):
#             continue
#         if not file.lower().endswith((".pdf", ".txt", ".docx")):
#             continue

#         try:
#             if file.endswith(".pdf"):
#                 with open(path, "rb") as f:
#                     reader = PyPDF2.PdfReader(f)
#                     text = "".join([page.extract_text() or "" for page in reader.pages])
#             elif file.endswith(".txt"):
#                 with open(path, "r", encoding="utf-8", errors='ignore') as f:
#                     text = f.read()
#             elif file.endswith(".docx"):
#                 doc = Document(path)
#                 text = "\n".join([p.text for p in doc.paragraphs])
#             else:
#                 continue

#             # Clean the text to remove problematic characters
#             text = clean_text(text)
            
#             if not text:
#                 print(f"Skipping empty document: {file}")
#                 continue

#             docs.append({
#                 "text": text,
#                 "metadata": {
#                     "filename": file,
#                     "path": path,
#                     "hash": file_hash(path),
#                     "timestamp": datetime.fromtimestamp(os.path.getmtime(path)).isoformat(),
#                 }
#             })
#         except Exception as e:
#             print(f"Error processing {file}: {str(e)}")
#             continue
    
#     return docs


# def chunk_texts(docs, chunk_size=1000, overlap=200):
#     """Split documents into overlapping text chunks."""
#     splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
#     chunks = []
#     for d in docs:
#         parts = splitter.create_documents([d["text"]], metadatas=[d["metadata"]])
#         chunks.extend(parts)
#     return chunks


# def build_pgvector_store(chunks):
#     """Embed and store document chunks in pgvector."""
#     embeddings = get_embeddings()
#     store = PGVector.from_documents(
#         documents=chunks,
#         embedding=embeddings,
#         connection_string=CONNECTION_STRING,
#         collection_name=rag_config.document_collection_name,
#     )
#     return store


# def load_existing_pgvector_store():
#     """Connect to existing pgvector collection."""
#     embeddings = get_embeddings()
#     logger.info(f"Loading existing pgvector store: {rag_config.document_collection_name}")
#     return PGVector(
#         connection_string=CONNECTION_STRING,
#         embedding_function=embeddings,
#         collection_name=rag_config.document_collection_name,
#     )


# def get_stored_hashes(store):
#     """Retrieve already-indexed file hashes."""
#     try:
#         with store._conn.cursor() as cur:
#             cur.execute(f"SELECT metadata->>'hash' FROM {rag_config.document_collection_name};")
#             return {r[0] for r in cur.fetchall() if r[0]}
#     except Exception:
#         return set()


# def update_index(force=False):
#     """Re-index new or changed files."""
#     if not force and not os.path.exists(rag_config.document_collection_path):
#         print(f"Folder {rag_config.document_collection_path} not found")
#         return
    
#     existing_store = load_existing_pgvector_store()
#     existing_hashes = get_stored_hashes(existing_store)

#     new_chunks = []
#     for doc in load_documents(rag_config.document_collection_path):
#         if doc["metadata"]["hash"] not in existing_hashes:
#             new_chunks.extend(chunk_texts([doc]))

#     if new_chunks:
#         print(f"Indexing {len(new_chunks)} new/updated chunks...")
#         build_pgvector_store(new_chunks)
#     else:
#         print("No new or changed documents.")


# def format_docs(docs):
#     """Format retrieved documents for the prompt."""
#     return "\n\n".join(doc.page_content for doc in docs)

# #main function to create the QA chain
# def make_qa_chain(skip_update=True):
#     """Create a Retrieval-Augmented QA chain using LCEL.
    
#     Args:
#         skip_update: If True, skip checking for new documents (faster).
#         config: RAGConfig instance or None to use global config
#         **kwargs: Override individual config parameters
#     """
#     # Use provided config or global config
#     # if config is None:
#     #     config = rag_config
    
#     # Override with any kwargs
#     #if kwargs:
#         # # Create a copy to avoid modifying the original
#         # config = RAGConfig()
#         # config.model_name = kwargs.get('model_name', rag_config.model_name)
#         # config.model_type = kwargs.get('model_type', rag_config.model_type)
#         # config.system_prompt = kwargs.get('system_prompt', rag_config.system_prompt)
#         # config.document_collection = kwargs.get('document_collection', rag_config.document_collection)
#     if not skip_update:
#      update_index()

#     store = load_existing_pgvector_store()
#     logger.info(f"Loaded PGVector store: {rag_config.document_collection_name}")
#     retriever = store.as_retriever(search_kwargs={"k": 5})
#     logger.info("Retriever created successfully")
#     # Create LLM based on model type
#     if rag_config.model_type.lower() == "openai":
#         logger.info("Using OpenAI LLM")
#         logger.info(f"model_name: {rag_config.model_name}")
#         llm = ChatOpenAI(model=rag_config.model_name, temperature=0, api_key=rag_config.openai_api_key)
#     elif rag_config.model_type.lower() == "local":
#         try:
#             # For local models, you can use Ollama or other local LLM providers
#             from langchain_community.llms import Ollama
#             llm = Ollama(model=rag_config.model_name, temperature=0)
#             print(f"Using local Ollama model: {rag_config.model_name}")
#         except ImportError: 
#             print("Ollama not available, falling back to OpenAI")
#     elif rag_config.model_type.lower() == "huggingface":
#         try:
#             # For HuggingFace models
#             from langchain_huggingface import HuggingFacePipeline
#             llm = HuggingFacePipeline.from_model_id(
#                 model_id=rag_config.model_name,
#                 task="text-generation",
#                 model_kwargs={"temperature": 0, "max_length": 512}
#             )
#             print(f"Using HuggingFace model: {rag_config.model_name}")
#         except ImportError:
#             print("HuggingFace not available, falling back to OpenAI")
#     else:
#         print(f"Unknown model type '{rag_config.model_type}', defaulting to OpenAI")
#         llm = ChatOpenAI(model=rag_config.model_name, temperature=0, api_key=rag_config.openai_api_key)

#     # Create prompt template
#     template = """{system_prompt}

# {context}

# Question: {input}

# Answer:"""
    
#     prompt = ChatPromptTemplate.from_template(template)

#     # Build the chain using LCEL (LangChain Expression Language)
#     def get_input(x):
#         """Extract input from dict or return string as-is."""
#         if isinstance(x, dict):
#             return x.get("input", x.get("question", ""))
#         return x

#     # Create the base chain
#     base_chain = (
#         {
#             # "context": lambda x: format_docs(retriever.invoke(get_input(x))),
#             "context": lambda x: format_docs(retriever.invoke(get_input(x))),
#             "input": get_input,
#             "system_prompt": lambda x: rag_config.system_prompt
#         }
#         | prompt
#         | llm
#         | StrOutputParser()
#     )
    
#     # Wrap to return dictionary format for compatibility
#     def wrapped_chain(x):
#         result = base_chain.invoke(x)
#         return {"output_text": result, "answer": result}
    
#     # Create a runnable wrapper
#     qa_chain = RunnableLambda(wrapped_chain)
    
#     return qa_chain